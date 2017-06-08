from docx import Document
from docx.shared import Inches
import PIL
import PIL.Image
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
from docx import Document
import re
import os
import sys
import pytesseract
from PIL import Image, ImageFont, ImageOps, ImageDraw, ImageEnhance, ImageFilter


def ImagetoDoc(filename):
    document = Document()
    name = raw_input('Enter name of the file which will save after conversion: ')+'.docx'
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(filename, width = Inches(6.0))
    document.save(name)
def ImagetoPDF(filename):
    name = raw_input('Enter name of the file which will save after conversion: ')+'.pdf'
    im = PIL.Image.open(filename)
    newfilename = name
    PIL.Image.Image.save(im, newfilename, "PDF", resolution = 100.0)
def PDFtoDoc(filename):
    document = Document()
    name = raw_input('Enter name of the file which will save after conversion: ')+'.docx'
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(filename, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)

    text = retstr.getvalue()
    myfile = text
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',r' ', myfile) # remove all non-XML-compatible characters
    illegal_xml_re = re.compile(u'[\x00-\x08\x0b-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufdd0-\ufddf\ufffe-\uffff]')
    myfile= illegal_xml_re.sub('', myfile)
    document.add_paragraph(myfile)
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('q.jpg', width = Inches(6.0))
    document.save(name)
    fp.close()
    device.close()
    retstr.close()
def PDFToText(filename):
    document = Document()
    fp = file(filename, 'rb')
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    # Create a PDF interpreter object.
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    # Process each page contained in the document.

    for page in PDFPage.get_pages(fp):
        interpreter.process_page(page)
        data =  retstr.getvalue()
    name = raw_input('Enter name of the file which will save after conversion: ')+'.txt'  # Name of text file coerced with +.txt
    fi = open(name,'a')
    fi.write(data)
    fi.close()
def DocToText(filename):
    document=Document(filename)
    filename=filename.strip('.docx') #not able to remove .docx
    name = raw_input('Enter name of the file which will save after conversion: ')+'.txt'  # Name of text file coerced with +.txt
    fi = open(name,'a')
    f=open(filename+".txt","wb")
    for para in document.paragraphs:
        f.write(para.text)
        f.close()
    f = para.text
    fi.write(f)
    fi.close()
def TextToDoc(filename):
    document = Document()
    myfile = open(filename).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
    p = document.add_paragraph(myfile)
    name = raw_input('Enter name of the file which will save after conversion: ')+'.docx'
    document.save(name)

# Image file to text file conversion using Pytesseract module using Tesseract-OCR
def image2text(file):
    image = Image.open(file) # the second one
    '''im = im.filter(ImageFilter.MedianFilter())
    enhancer = ImageEnhance.Contrast(im)
    im = enhancer.enhance(2)
    im = im.convert('1')
    im.save('sample/temp.jpg')'''
    text = pytesseract.image_to_string(image,lang='eng').encode('cp850','replace').decode('cp850')
    textfile = open("output/output.txt","w")
    textfile.write(text)

    # Text file to image file conversion
def text2image(text_path, font_path=None):
    """Convert text file to a grayscale image with black characters on a white background.

    arguments:
    text_path - the content of this file will be converted to an image
    font_path - path to a font file (for example impact.ttf)
    """
    grayscale = 'L'
    # parse the file into lines
    with open(text_path) as text_file:  # can throw FileNotFoundError
        lines = tuple(l.rstrip() for l in text_file.readlines())

    # choose a font (you can see more detail in my library on github)
    large_font = 20  # get better resolution with larger size
    font_path = font_path or 'cour.ttf'  # Courier New. works in windows. linux may need more explicit path
    try:
        font = PIL.ImageFont.truetype(font_path, size=large_font)
    except IOError:
        font = PIL.ImageFont.load_default()
        print('Could not use chosen font. Using default.')

    # make the background image based on the combination of font and lines
    pt2px = lambda pt: int(round(pt * 96.0 / 72))  # convert points to pixels
    max_width_line = max(lines, key=lambda s: font.getsize(s)[0])
    # max height is adjusted down because it's too large visually for spacing
    test_string = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    max_height = pt2px(font.getsize(test_string)[1])
    max_width = pt2px(font.getsize(max_width_line)[0])
    height = max_height * len(lines)  # perfect or a little oversized
    width = int(round(max_width + 40))  # a little oversized
    image = PIL.Image.new(grayscale, (width, height), color=255)
    draw = PIL.ImageDraw.Draw(image)

    # draw each line of text
    vertical_position = 5
    horizontal_position = 5
    line_spacing = int(round(max_height * 0.8))  # reduced spacing seems better
    for line in lines:
        draw.text((horizontal_position, vertical_position),
                  line, fill=0, font=font)
        vertical_position += line_spacing
    # crop the text
    c_box = PIL.ImageOps.invert(image).getbbox()
    image = image.crop(c_box)
    image.save('output/result.jpg')

#doc file to image file conversion using doc2text + text2image file conversion
def doc2image(file):
    document=Document(filename)
    filename=filename.strip('.docx') #not able to remove .docx
    name = raw_input('Enter name of the file which will save after conversion: ')+'.txt'  # Name of text file coerced with +.txt
    fi = open(name,'a')
    f=open(filename+".txt","wb")
    for para in document.paragraphs:
        f.write(para.text)
        f.close()
    f = para.text
    fi.write(f)
    fi.close()
    text2image(fi+".txt")

def printhello():
	print("Hellow world!")

def convertfile(file,qurytype):
    print(querytype)
    if querytype=="PDF to Text":
        pdf2text(file)
    elif querytype=="PDF to Doc":
        pdf2doc(file)
    elif querytype=="Text to Doc":
        text2doc(file)
    elif querytype=="Doc to Text":
        doc2text(file)
    elif querytype=="Image to Text":
        image2text(file)
    elif querytype=="Text to Image":
        text2image(file)
    elif querytype=="Image to Doc":
        image2doc(file)
    elif querytype=="Doc to Image":
        doc2image(file)
    elif querytype=="Image to PDF":
        image2pdf(file)
    
